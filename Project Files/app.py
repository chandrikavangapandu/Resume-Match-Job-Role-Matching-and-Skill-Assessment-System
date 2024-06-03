import streamlit as st
import json
import os
import numpy as np
from docx import Document
import pdfplumber
import matplotlib.colors as mcolors
from operator import index
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.probability import FreqDist
import string
from dotenv import load_dotenv
import docx
from PIL import Image
from openai import OpenAI
import os
import nltk


load_dotenv()

@st.cache_resource
def download_resource():
    nltk.download("punkt")
    nltk.download("stopwords")

session_state = st.session_state
if "user_index" not in st.session_state:
    st.session_state["user_index"] = 0


download_resource()

def signup(json_file_path="data.json"):
    st.title("Signup Page")
    with st.form("signup_form"):
        st.write("Fill in the details below to create an account:")
        name = st.text_input("Name:")
        username = st.text_input("username:")
        age = st.number_input("Age:", min_value=0, max_value=120)
        sex = st.radio("Sex:", ("Male", "Female", "Other"))
        password = st.text_input("Password:", type="password")
        confirm_password = st.text_input("Confirm Password:", type="password")

        if st.form_submit_button("Signup"):
            if password == confirm_password:
                user = create_account(
                    name, username, age, sex, password, json_file_path
                )
                session_state["logged_in"] = True
                session_state["user_info"] = user
            else:
                st.error("Passwords do not match. Please try again.")


def check_login(username, password, json_file_path="data.json"):
    try:
        with open(json_file_path, "r") as json_file:
            data = json.load(json_file)

        for user in data["users"]:
            if user["username"] == username and user["password"] == password:
                session_state["logged_in"] = True
                session_state["user_info"] = user
                st.success("Login successful!")
                return user

        st.error("Invalid credentials. Please try again.")
        return None
    except Exception as e:
        st.error(f"Error checking login: {e}")
        return None


def initialize_database(json_file_path="data.json"):
    try:
        # Check if JSON file exists
        if not os.path.exists(json_file_path):
            # Create an empty JSON structure
            data = {"users": []}
            with open(json_file_path, "w") as json_file:
                json.dump(data, json_file)
    except Exception as e:
        print(f"Error initializing database: {e}")


def create_account(name, username, age, sex, password, json_file_path="data.json"):
    try:
        # Check if the JSON file exists or is empty
        if not os.path.exists(json_file_path) or os.stat(json_file_path).st_size == 0:
            data = {"users": []}
        else:
            with open(json_file_path, "r") as json_file:
                data = json.load(json_file)
        username = username.lower()
        

        # Append new user data to the JSON structure
        user_info = {
            "name": name,
            "username": username,
            "age": age,
            "sex": sex,
            "password": password,
            "resume": None,
            "job_description": None,
            "job_applied": None,
            "score": "0",
            "questions": None,
        }
        data["users"].append(user_info)

        # Save the updated data to JSON
        with open(json_file_path, "w") as json_file:
            json.dump(data, json_file, indent=4)

        st.success("Account created successfully! You can now login.")
        return user_info
    except json.JSONDecodeError as e:
        st.error(f"Error decoding JSON: {e}")
        return None
    except Exception as e:
        st.error(f"Error creating account: {e}")
        return None


def login(json_file_path="data.json"):
    st.title("Login Page")
    username = st.text_input("Username:")
    password = st.text_input("Password:", type="password")

    login_button = st.button("Login")

    if login_button:
        username = username.lower()
        user = check_login(username, password, json_file_path)
        if user is not None:
            session_state["logged_in"] = True
            session_state["user_info"] = user
        else:
            st.error("Invalid credentials. Please try again.")


def get_user_info(username, json_file_path="data.json"):
    try:
        with open(json_file_path, "r") as json_file:
            data = json.load(json_file)
            for user in data["users"]:
                if user["username"] == username:
                    return user
        return None
    except Exception as e:
        st.error(f"Error getting user information: {e}")
        return None


def render_dashboard(user_info, json_file_path="data.json"):
    try:
        st.title(f"Welcome to the Dashboard, {user_info['name']}!")
        st.subheader("User Information:")
        st.write(f"Name: {user_info['name']}")
        st.write(f"Sex: {user_info['sex']}")
        st.write(f"Age: {user_info['age']}")
        
        if user_info["resume"] is not None:
            st.subheader("Resume:")
            st.write(user_info["resume"])
            
        if user_info["job_description"] is not None:
            st.subheader("Job Applied:")
            st.write(user_info["job_description"])
        
    except Exception as e:
        st.error(f"Error rendering dashboard: {e}")


def extract_text(file) -> str:
    if isinstance(file, str):
        file_extension = os.path.splitext(file)[1].lower()
    else:
        file_extension = os.path.splitext(file.name)[1].lower()

    if file_extension == ".pdf":
        if isinstance(file, str):
            with pdfplumber.open(file) as pdf:
                text = "\n".join(
                    page.extract_text() for page in pdf.pages if page.extract_text()
                )
        else:
            with pdfplumber.open(file) as pdf:
                text = "\n".join(
                    page.extract_text() for page in pdf.pages if page.extract_text()
                )
    elif file_extension == ".docx":
        if isinstance(file, str):
            doc = docx.Document(file)
        else:
            doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        if isinstance(file, str):
            with open(file, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        else:
            with file as f:
                text = f.read()
    return text


def process_text(text):
    tokens = word_tokenize(text)
    tokens = [word.lower() for word in tokens if word.isalpha()]
    stop_words = set(stopwords.words("english"))
    tokens = [word for word in tokens if word not in stop_words]
    return tokens


def calculate_score(job_description_tokens, resume_tokens):
    job_description_freq = FreqDist(job_description_tokens)
    total_tokens_in_job_description = len(job_description_tokens)
    resume_tokens = list(set(resume_tokens))
    score = sum(job_description_freq[token] for token in resume_tokens)
    score_percentage = (score / total_tokens_in_job_description) * 100
    return score_percentage


def extract_keywords_from_resume(resume_text):
    resume_text = resume_text.lower()
    resume_tokens = word_tokenize(resume_text)
    resume_tokens = [
        token for token in resume_tokens if token not in string.punctuation
    ]
    stop_words = set(stopwords.words("english"))
    resume_tokens = [token for token in resume_tokens if token not in stop_words]
    processed_resume_text = " ".join(resume_tokens)
    client = OpenAI(
        api_key=os.environ.get("OPENAI_API_KEY"),
    )
    prompt = f"Extract top most important skill keywords from the given resume text. Separate the skills with comma.\n{processed_resume_text}\nKeywords:"
    messages = [{"role": "system", "content": prompt}]
    response = client.chat.completions.create(
        messages=messages,
        model="gpt-3.5-turbo",
    )
    return response.choices[0].message.content.split(",")
    
def suggested_skills(resume_text):
    client = OpenAI(
        api_key=os.environ.get("OPENAI_API_KEY"),
    )
    prompt = f"""Extract top skills that are not mentioned in the resume but are relevant to tech jobs. 
    \n{resume_text}\nMissing Keywords:"""
    messages = [{"role": "system", "content": prompt}]
    response = client.chat.completions.create(
        messages=messages,
        model="gpt-3.5-turbo",
    )
    return response.choices[0].message.content.split(",")


def resume_score(resume_text):
    client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
    prompt = f"""Analyze the provided resume text to generate a score reflecting the candidate's skills and experience, rated out of 100. Only return the score as an integer. Do not include any other information in the response. Only return the score as an integer.


Resume Text:
\n{resume_text}"""
    messages = [{"role": "system", "content": prompt}]
    response = client.chat.completions.create(
        messages=messages,
        model="gpt-3.5-turbo",
    )
    return int(response.choices[0].message.content)



def suggestions(resume_text):
    client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
    prompt = f"""Analyze the provided resume text. Highlight the candidate's skills and experience, and suggest improvements to make the resume more appealing to potential employers.
    
    Task:
    1. Analyze the resume text to identify the candidate's skills and experience and provide a brief summary.
    2. Suggest improvements to make the resume more appealing to potential employers.
    For each skill or experience mentioned in the resume, provide a brief explanation or example to demonstrate the candidate's proficiency or relevance to the job.
    Resume Text:
    \n{resume_text}"""
    messages = [{"role": "system", "content": prompt}]
    response = client.chat.completions.create(
        messages=messages,
        model="gpt-3.5-turbo",
    )
    return response.choices[0].message.content

def evaluate_interview(resume_text, job_description_text, candidate_name, questions, responses):
    client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
    question_responses = "\n".join(
        f"Q: {question}\nA: {response}\n" for question, response in zip(questions, responses)
    )
    
    
    prompt = f"""Evaluate the candidate's responses to the interview questions based on the provided resume and job description. Rate the candidate's performance on a scale of 0 to 100, with 0 being the lowest and 100 being the highest. Provide feedback on the candidate's strengths and areas for improvement. 
    
    Format: Score: Score \n Feedback: Feedback (Detailed feedback on the candidate's performance and areas for improvement)
    
    Resume: {resume_text}
    
    Job Description: {job_description_text}
    
    Candidate Name: {candidate_name}
    
    Interview Questions:
    {question_responses}
    
    
    """
    messages = [{"role": "system", "content": prompt}]
    response = client.chat.completions.create(
        messages=messages,
        model="gpt-3.5-turbo",
    )
    result = response.choices[0].message.content
    score = result.split("Score: ")[1].split("Feedback:")[0].strip()
    feedback = result.split("Feedback: ")[1].strip()
    return int(score), feedback

def generate_question(
    resume_text,
    job_description_text,
    candidate_name,
    previous_response=None,
    previous_question=None,
):
    prompt = f"Resume Text: {resume_text}\nJob Description: {job_description_text}\nCandidate Name: {candidate_name}\n"
    if previous_response and previous_question:
        prompt += f"Previous Response: {previous_response}\nPrevious Question: {previous_question}\n"
    client = OpenAI(
        api_key=os.environ.get("OPENAI_API_KEY"),
    )
    messages = [
        {"role": "system", "content": "You are the interviewer."},
        {
            "role": "system",
            "content": "You are interviewing a candidate. Ask a question based on the resume and job description. If the candidate has already answered a question, you can ask a follow-up question based on their response.",
        },
        {"role": "user", "content": prompt},
    ]
    response = client.chat.completions.create(
        messages=messages,
        model="gpt-3.5-turbo",
    )
    return response.choices[0].message.content


def main(json_file_path="data.json"):
    st.sidebar.title("Resume Screening system")
    page = st.sidebar.selectbox(
        "Go to",
        (
            "Signup/Login",
            "Dashboard",
            "Resume Analysis",
            "Apply for a job",
            "Generate Questions",
            "Candidate Evaluation",
            "Logout",
        ),
        key="GET YOUR RESUME ANALYZED AND COMPARED",
    )

    if page == "Signup/Login":
        st.title("Signup/Login Page")
        login_or_signup = st.radio(
            "Select an option", ("Login", "Signup"), key="login_signup"
        )
        if login_or_signup == "Login":
            login(json_file_path)
        else:
            signup(json_file_path)
            
    elif page == "Dashboard":
        if session_state.get("logged_in"):
            user_info = session_state["user_info"]
            render_dashboard(user_info, json_file_path)
        else:
            st.warning("Please login/signup to view the dashboard.")

    elif page == "Resume Analysis":
        if session_state.get("logged_in"):
            st.title("Upload resume for analysis")
            uploaded_file = st.file_uploader("Choose a file", type=None)
            if uploaded_file is not None:
                resume_text = extract_text(uploaded_file)
                st.write("File name: ", uploaded_file.name)
                st.success("File uploaded successfully!")
                # st.image(Image.open("Images/logo.png"), use_column_width=True)
                with open(json_file_path, "r+") as json_file:
                    data = json.load(json_file)
                    user_index = next(
                        (
                            i
                            for i, user in enumerate(data["users"])
                            if user["username"]
                            == session_state["user_info"]["username"]
                        ),
                        None,
                    )
                    if user_index is not None:
                        user_info = data["users"][user_index]
                        user_info["resume"] = resume_text
                        session_state["user_info"] = user_info
                        json_file.seek(0)
                        json.dump(data, json_file, indent=4)
                        json_file.truncate()
                    else:
                        st.error("User not found.")
                resume_keywords = extract_keywords_from_resume(resume_text)
                st.subheader("Skills of the candidate:")
                # use st.multiselect to display the keywords
                skills = st.multiselect("Candidate Skills", resume_keywords, resume_keywords)
                score  = resume_score(resume_text)
                suggestions_improvements = suggestions(resume_text)
                st.subheader("Resume Score")
                if score >= 60:
                    st.success(
                        f"Congratulations! Your resume score is {score}. Your resume is well-matched for technical roles"
                    )
                elif score >= 20:
                    st.warning(
                        f"Your resume score is {score}. Your resume is not well-matched for technical roles. Consider improvements."
                    )
                else:
                    st.error(
                        f"Your resume score is {score}. Your resume is not well-matched for technical roles. Consider significant improvements."
                    )
                percentage_score = score / 100
                percentage_remainder = 1 - percentage_score

                # Create a Plotly figure for the pie chart
                fig = go.Figure(
                data=[
                    go.Pie(
                        labels=["Matched", "Unmatched"],
                        values=[percentage_score, percentage_remainder],
                        hole=0.3,
                        marker_colors=["rgba(0, 128, 0, 0.7)", "rgba(255, 0, 0, 0.7)"],
                    )
                ]
            )
                fig.update_layout(title_text="Resume Score")

                # Display the chart
                st.plotly_chart(fig)
                st.subheader("Skills and Experience:")
                st.write(suggestions_improvements)
                skills_not_mentioned = suggested_skills(resume_text)
                st.subheader("Top recommended skills:")
                skills = st.multiselect("Recommended Skills", skills_not_mentioned, skills_not_mentioned)
                st.subheader("How does your resume compare with other candidates?")
                scores = [
                    int(user["score"])
                    for user in data["users"]
                ]
                # Plot an interactive graph
                fig = go.Figure()
                fig.add_trace(
                    go.Histogram(
                        x=scores,
                        histnorm="percent",
                        marker_color="rgba(0, 0, 255, 0.7)",
                        opacity=0.75,
                    )
                )

                fig.update_layout(
                    title_text=f"Distribution of Scores",
                    xaxis_title="Resume Score",
                    yaxis_title="Percentage of Candidates",
                    bargap=0.05,
                )

                st.plotly_chart(fig)

        else:
            st.warning("Please login/signup to view the dashboard.")
    
    
    elif page == "Apply for a job":
        if session_state.get("logged_in"):
            st.title("Apply for a job")
            st.subheader("Select a role to you want to apply for:")
            
            BASE_DIR = "Data\\JobDesc\\"
            job_description = st.selectbox(
                "Select a role",
                [
                    "-Select-",
                    "Backend Developer",
                    "Billing cum Logistics Manager",
                    "Data Scientist",
                    "Director of Engineering",
                    "Global Industry Content Manager",
                    "HTML Developer",
                    "IT Project Manager",
                    "Lead Technical Program Manager",
                    "Primary English Teacher",
                    "Revenue Reporting Data Analyst",
                    "Senior Product Manager",
                    "Senior Software Developer",
                    "Web Developer",
                    "Web_dev_job",
                ],
                key="job_description",
            )
            if job_description and job_description != "-Select-":
                file_path = os.path.join(BASE_DIR, f"{job_description}.docx")
                job_description_text = extract_text(file_path)
                st.subheader("Job Description:")
                st.write(job_description_text)
                if st.button("Apply"):
                    with open(json_file_path, "r+") as json_file:
                        data = json.load(json_file)
                        user_index = next((i for i, user in enumerate(data["users"]) if user["username"] == session_state["user_info"]["username"]), None)
                        if user_index is not None:
                            user_info = data["users"][user_index]
                            user_info["job_description"] = job_description_text
                            user_info["job_applied"] = job_description
                            session_state["user_info"] = user_info
                            json_file.seek(0)
                            json.dump(data, json_file, indent=4)
                            json_file.truncate()
                        else:
                            st.error("User not found.")
                    st.success("Job application submitted successfully!")
        else:
            st.warning("Please login/signup to view the dashboard.")
    elif page == "Generate Questions":
        if session_state.get("logged_in"):
            user_info = session_state["user_info"]
            st.title("Give your interview")
            st.write("Answer the questions below to complete the interview.")
            with open(json_file_path, "r") as json_file:
                data = json.load(json_file)
                user_index = next(
                    (
                        i
                        for i, user in enumerate(data["users"])
                        if user["username"] == session_state["user_info"]["username"]
                    ),
                    None,
                )
                if (
                session_state["user_info"]["resume"] is None
                or session_state["user_info"]["job_description"] is None
            ):
                    st.warning(
                    "Please upload your resume and apply for a job to generate interview questions."
                )
                    return
            st.markdown("### Interview Questions")

            if "messages" not in st.session_state:
                st.session_state.messages = []

            if user_info["questions"] is None:
                previous_response = None
                previous_question = None
            else:
                previous_response = user_info["questions"][-1]["response"]
                previous_question = user_info["questions"][-1]["question"]
            if user_info["questions"] is not None and len(user_info["questions"]) > 0:
                for questions in user_info["questions"]:
                    st.chat_message("Interviewer", avatar="🤖").write(questions["question"])
                    st.chat_message("Applicant", avatar="👩‍🎨").write(questions["response"])
                
            question = generate_question(
                session_state["user_info"]["resume"],
                session_state["user_info"]["job_description"],
                session_state["user_info"]["name"],
                previous_response,
                previous_question,
            )
            with st.chat_message("Interviewer", avatar="🤖"):
                st.markdown(question)

            if prompt := st.chat_input("Enter your response here", key="response"):
                with st.chat_message("Applicant", avatar="👩‍🎨"):
                    st.markdown(prompt)

                with open(json_file_path, "r+") as json_file:
                    data = json.load(json_file)
                    user_index = next(
                        (
                            i
                            for i, user in enumerate(data["users"])
                            if user["username"] == session_state["user_info"]["username"]
                        ),
                        None,
                    )
                    if user_index is not None:
                        user_info = data["users"][user_index]
                        if user_info["questions"] is None:
                            user_info["questions"] = []
                        user_info["questions"].append(
                            {"question": question, "response": prompt}
                        )
                        session_state["user_info"] = user_info
                        json_file.seek(0)
                        json.dump(data, json_file, indent=4)
                        json_file.truncate()
                    else:
                        st.error("User not found.")

                st.rerun()

            if st.button("Finish Interview"):
                with open(json_file_path, "r+") as json_file:
                    data = json.load(json_file)
                    user_index = next(
                        (
                            i
                            for i, user in enumerate(data["users"])
                            if user["username"]
                            == session_state["user_info"]["username"]
                        ),
                        None,
                    )
                    if user_index is not None:
                        user_info = data["users"][user_index]
                        session_state["user_info"] = user_info
                        json_file.seek(0)
                        json.dump(data, json_file, indent=4)
                        json_file.truncate()
                    else:
                        st.error("User not found.")
                st.success("Interview questions completed successfully!")
                return
                
        else:
            st.warning("Please login/signup to give your interview.")
    
    elif page == "Candidate Evaluation":
        if session_state.get("logged_in"):
            user_info = session_state["user_info"]
            st.title("Candidate Evaluation based on Resume and Interview")
            st.write("Evaluate the candidate's responses to the interview questions.")
            questions = [user_info["questions"][i]["question"] for i in range(len(user_info["questions"]))]
            responses = [user_info["questions"][i]["response"] for i in range(len(user_info["questions"]))]
            score, feedback = evaluate_interview(
                user_info["resume"],
                user_info["job_description"],
                user_info["name"],
                questions,
                responses,
            )
            st.write(f"Score: {score}")
            percentage_score = score / 100
            percentage_remainder = 1 - percentage_score

                # Create a Plotly figure for the pie chart
            fig = go.Figure(
                data=[
                    go.Pie(
                        labels=["Matched", "Unmatched"],
                        values=[percentage_score, percentage_remainder],
                        hole=0.3,
                        marker_colors=["rgba(0, 128, 0, 0.7)", "rgba(255, 0, 0, 0.7)"],
                    )
                ]
            )
            fig.update_layout(title_text="Resume Score")

                # Display the chart
            st.plotly_chart(fig)
            
            for (question, response) in zip(questions, responses):
                st.markdown(f"##### Question:")
                st.write(question)
                st.markdown(f"##### Response:")
                st.write(response)
            st.subheader("Feedback:")
            st.markdown(f"{feedback}")
    elif page == "Logout":
        st.title("Logout")
        if st.button("Logout"):
            session_state["logged_in"] = False
            session_state["user_info"] = None
            st.success("You have been logged out successfully.")
        st.image("Images\logo.png", use_column_width=True)


if __name__ == "__main__":
    initialize_database()
    main()
